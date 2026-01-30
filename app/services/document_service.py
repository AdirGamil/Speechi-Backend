"""
Document generation service (Word only).

Builds Word (.docx) documents from meeting transcript and analysis.
PDF export is handled by pdf_service (WeasyPrint: HTML → PDF).
No AI usage; consumes transcript and AnalysisResult only.

Features:
- Language-aware section headings
- RTL support for Hebrew and Arabic
- Clean, professional typography
"""

import os
import tempfile
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.models.schemas import ActionItem, AnalysisResult
from app.utils import file_utils
from app.utils.document_labels import get_labels, is_rtl, DocumentLabels


def _format_action_item(item: ActionItem, labels: DocumentLabels) -> str:
    """Format a single action item as 'description (Owner: X)' or '(Owner: Unassigned)'."""
    owner = item.owner.strip() if item.owner else None
    if owner:
        return f"{item.description} ({labels['owner']}: {owner})"
    return f"{item.description} ({labels['owner']}: {labels['unassigned']})"


# ============================================
# Word Document Generation
# ============================================

def _set_rtl_paragraph(paragraph) -> None:
    """Set RTL direction for a Word paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)


def generate_word_document(analysis: AnalysisResult, language: str = "en") -> str:
    """
    Build a Word document from analysis with language-aware headings.

    Includes both clean/translated transcript and original transcript.
    For long recordings, the clean transcript may be condensed while the
    original transcript is always complete.

    Args:
        analysis: Structured analysis (summary, participants, decisions,
            action_items, translated_transcript, raw_transcript).
        language: Output language code (en, he, fr, es, ar). Default en.

    Returns:
        Absolute path to the generated .docx file. Caller is responsible for cleanup.

    Raises:
        OSError: If the temp file cannot be created or written.
    """
    labels = get_labels(language)
    rtl = is_rtl(language)
    doc = Document()

    # Set document language for RTL if needed
    if rtl:
        # Set default paragraph direction
        style = doc.styles['Normal']
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Title
    title = doc.add_heading(labels["title"], level=0)
    if rtl:
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_rtl_paragraph(title)

    # Summary section
    heading = doc.add_heading(labels["summary"], level=1)
    if rtl:
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_rtl_paragraph(heading)
    
    summary_text = (analysis.summary or "").strip()
    para = doc.add_paragraph(summary_text if summary_text else f"({labels['none']})")
    if rtl:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_rtl_paragraph(para)

    # Clean Transcript section (translated/condensed)
    transcript_label = labels["clean_transcript"]
    if getattr(analysis, "is_condensed", False):
        transcript_label += f" {labels['condensed_note']}"
    
    heading = doc.add_heading(transcript_label, level=1)
    if rtl:
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_rtl_paragraph(heading)
    
    transcript_text = (analysis.translated_transcript or "").strip()
    para = doc.add_paragraph(transcript_text if transcript_text else f"({labels['none']})")
    if rtl:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_rtl_paragraph(para)

    # Original Transcript section (raw Whisper output)
    raw_transcript = getattr(analysis, "raw_transcript", "") or ""
    if raw_transcript.strip():
        heading = doc.add_heading(labels["original_transcript"], level=1)
        if rtl:
            heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_rtl_paragraph(heading)
        
        para = doc.add_paragraph(raw_transcript.strip())
        if rtl:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_rtl_paragraph(para)

    # Participants section
    heading = doc.add_heading(labels["participants"], level=1)
    if rtl:
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_rtl_paragraph(heading)
    
    if not analysis.participants:
        para = doc.add_paragraph(labels["none"])
        if rtl:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_rtl_paragraph(para)
    else:
        for p in analysis.participants:
            para = doc.add_paragraph(p, style="List Bullet")
            if rtl:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _set_rtl_paragraph(para)

    # Decisions section
    heading = doc.add_heading(labels["decisions"], level=1)
    if rtl:
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_rtl_paragraph(heading)
    
    if not analysis.decisions:
        para = doc.add_paragraph(labels["none"])
        if rtl:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_rtl_paragraph(para)
    else:
        for d in analysis.decisions:
            para = doc.add_paragraph(d, style="List Bullet")
            if rtl:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _set_rtl_paragraph(para)

    # Action Items section
    heading = doc.add_heading(labels["actions"], level=1)
    if rtl:
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_rtl_paragraph(heading)
    
    if not analysis.action_items:
        para = doc.add_paragraph(labels["none"])
        if rtl:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_rtl_paragraph(para)
    else:
        for item in analysis.action_items:
            para = doc.add_paragraph(_format_action_item(item, labels), style="List Bullet")
            if rtl:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _set_rtl_paragraph(para)

    # Footer
    doc.add_paragraph()  # Spacer
    footer = doc.add_paragraph(f"{labels['generated_by']} • {datetime.now().strftime('%Y-%m-%d')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # Save to temp file
    fd, path = tempfile.mkstemp(suffix=".docx")
    try:
        os.close(fd)
        doc.save(path)
    except Exception:
        file_utils.delete_temp_file(path)
        raise
    return path


# ============================================
# Legacy function (backward compatibility)
# ============================================

def generate_document(analysis: AnalysisResult, language: str = "en") -> str:
    """
    Build a Word document from analysis (backward-compatible wrapper).
    
    Delegates to generate_word_document.
    """
    return generate_word_document(analysis, language)
